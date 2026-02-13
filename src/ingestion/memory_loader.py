"""In-memory document loader for user uploads and URLs."""

import io
import tempfile
from typing import List, BinaryIO, Optional
from pathlib import Path
from langchain_core.documents import Document


class InMemoryDocumentLoader:
    """Load documents from memory (uploads, URLs) without persisting to disk."""

    def __init__(self):
        """Initialize in-memory loader."""
        pass

    def load_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        file_type: str = "pdf"
    ) -> List[Document]:
        """Load document from bytes.

        Args:
            file_bytes: File content as bytes
            filename: Original filename
            file_type: File type (pdf, docx, txt, png, jpg, jpeg, webp)

        Returns:
            List of Document objects
        """
        if file_type.lower() == "pdf":
            return self._load_pdf_from_bytes(file_bytes, filename)
        elif file_type.lower() in ["docx", "doc"]:
            return self._load_docx_from_bytes(file_bytes, filename)
        elif file_type.lower() == "txt":
            return self._load_txt_from_bytes(file_bytes, filename)
        elif file_type.lower() in ["png", "jpg", "jpeg", "webp", "gif", "bmp"]:
            return self._load_image_from_bytes(file_bytes, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _load_pdf_from_bytes(self, pdf_bytes: bytes, filename: str) -> List[Document]:
        """Load PDF from bytes using PyMuPDF."""
        try:
            import pymupdf

            # Open PDF from bytes
            doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
            documents = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "filename": filename,
                            "page": page_num,
                            "source": "upload",
                            "total_pages": len(doc),
                        }
                    ))

            doc.close()
            return documents

        except ImportError:
            # Fallback to pypdf
            return self._load_pdf_with_pypdf(pdf_bytes, filename)

    def _load_pdf_with_pypdf(self, pdf_bytes: bytes, filename: str) -> List[Document]:
        """Fallback: Load PDF using pypdf."""
        from pypdf import PdfReader

        pdf_file = io.BytesIO(pdf_bytes)
        reader = PdfReader(pdf_file)
        documents = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()

            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "filename": filename,
                        "page": page_num,
                        "source": "upload",
                        "total_pages": len(reader.pages),
                    }
                ))

        return documents

    def _load_docx_from_bytes(self, docx_bytes: bytes, filename: str) -> List[Document]:
        """Load DOCX from bytes."""
        from docx import Document as DocxDocument

        docx_file = io.BytesIO(docx_bytes)
        doc = DocxDocument(docx_file)

        # Combine all paragraphs
        text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        return [Document(
            page_content=text,
            metadata={
                "filename": filename,
                "page": 0,
                "source": "upload",
                "total_pages": 1,
            }
        )]

    def _load_txt_from_bytes(self, txt_bytes: bytes, filename: str) -> List[Document]:
        """Load plain text from bytes."""
        text = txt_bytes.decode('utf-8', errors='ignore')

        return [Document(
            page_content=text,
            metadata={
                "filename": filename,
                "page": 0,
                "source": "upload",
                "total_pages": 1,
            }
        )]

    def _load_image_from_bytes(self, image_bytes: bytes, filename: str) -> List[Document]:
        """Load image from bytes using OCR."""
        try:
            from src.extraction.image_ocr import ImageOCRProcessor

            # Initialize OCR processor
            ocr_processor = ImageOCRProcessor()

            # Extract text from image
            text = ocr_processor.extract_text_from_image_bytes(image_bytes)

            return [Document(
                page_content=text,
                metadata={
                    "filename": filename,
                    "page": 0,
                    "source": "upload",
                    "type": "image",
                    "total_pages": 1,
                }
            )]
        except Exception as e:
            # If OCR fails, return minimal document
            return [Document(
                page_content=f"[Image: {filename}] - OCR processing failed: {str(e)}",
                metadata={
                    "filename": filename,
                    "page": 0,
                    "source": "upload",
                    "type": "image",
                    "total_pages": 1,
                }
            )]

    def load_from_upload(self, uploaded_file) -> List[Document]:
        """Load document from Streamlit uploaded file.

        Args:
            uploaded_file: Streamlit UploadedFile object

        Returns:
            List of Document objects
        """
        file_bytes = uploaded_file.read()
        filename = uploaded_file.name
        file_type = Path(filename).suffix[1:]  # Remove the dot

        return self.load_from_bytes(file_bytes, filename, file_type)

    def load_from_url(self, url: str) -> List[Document]:
        """Load document from URL.

        Args:
            url: URL to fetch document from

        Returns:
            List of Document objects
        """
        import requests

        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            # Determine file type from URL or Content-Type
            content_type = response.headers.get('Content-Type', '')

            if 'pdf' in content_type or url.lower().endswith('.pdf'):
                file_type = 'pdf'
            elif 'word' in content_type or url.lower().endswith(('.docx', '.doc')):
                file_type = 'docx'
            else:
                file_type = 'txt'

            filename = Path(url).name or "downloaded_document"

            return self.load_from_bytes(response.content, filename, file_type)

        except Exception as e:
            raise RuntimeError(f"Error loading from URL: {e}")
